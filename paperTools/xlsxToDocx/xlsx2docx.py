import tkinter as tk
from tkinter import ttk, scrolledtext
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt

# 打包：
# pyinstaller --onefile --noconsole xlsx2docx.py

# 创建主窗口
root = tk.Tk()
root.title("Excel to Word Converter")
root.geometry("560x680")
root.configure(bg="#f0f0f0")  # 设置背景颜色

# 定义默认值
default_excel_file = "input.xlsx"
default_docx_file = "output.docx"
default_r_title = 1
default_r_start = 2
default_r_end = 4
default_c_start = 'A'
default_c_end = 'D'

# 定义变量
excel_file_var = tk.StringVar(value=default_excel_file)
docx_file_var = tk.StringVar(value=default_docx_file)
r_title_var = tk.IntVar(value=default_r_title)
r_start_var = tk.IntVar(value=default_r_start)
r_end_var = tk.IntVar(value=default_r_end)
c_start_var = tk.StringVar(value=default_c_start)
c_end_var = tk.StringVar(value=default_c_end)

# 选项变量
bold_header_vars = []
bold_content_vars = []
colon_vars = []
newline_vars = []
extra_newline_vars = []

# 当前 options_frame 覆盖的范围
current_options_range = ("A", "D")

# 创建输入框
def create_input_frame(parent):
    frame = tk.Frame(parent, bg="#f0f0f0")
    tk.Label(frame, text="导入 Excel 名:", bg="#f0f0f0").grid(row=0, column=0, padx=5, pady=5, sticky="e")
    tk.Entry(frame, textvariable=excel_file_var, width=30).grid(row=0, column=1, columnspan=3, padx=5, pady=5, sticky="w")
    tk.Label(frame, text="生成 Word 名:", bg="#f0f0f0").grid(row=1, column=0, padx=5, pady=5, sticky="e")
    tk.Entry(frame, textvariable=docx_file_var, width=30).grid(row=1, column=1, columnspan=3, padx=5, pady=5, sticky="w")
    tk.Label(frame, text="题目行号:", bg="#f0f0f0").grid(row=2, column=0, padx=5, pady=5, sticky="e")
    tk.Entry(frame, textvariable=r_title_var, width=10).grid(row=2, column=1, padx=5, pady=5, sticky="w")
    tk.Label(frame, text="开始行号:", bg="#f0f0f0").grid(row=3, column=0, padx=5, pady=5, sticky="e")
    tk.Entry(frame, textvariable=r_start_var, width=10).grid(row=3, column=1, padx=5, pady=5, sticky="w")
    tk.Label(frame, text="结束行号:", bg="#f0f0f0").grid(row=3, column=2, padx=5, pady=5, sticky="e")
    tk.Entry(frame, textvariable=r_end_var, width=10).grid(row=3, column=3, padx=5, pady=5, sticky="w")
    tk.Label(frame, text="开始列号:", bg="#f0f0f0").grid(row=4, column=0, padx=5, pady=5, sticky="e")
    tk.Entry(frame, textvariable=c_start_var, width=10).grid(row=4, column=1, padx=5, pady=5, sticky="w")
    tk.Label(frame, text="结束列号:", bg="#f0f0f0").grid(row=4, column=2, padx=5, pady=5, sticky="e")
    tk.Entry(frame, textvariable=c_end_var, width=10).grid(row=4, column=3, padx=5, pady=5, sticky="w")

    # 添加“段间是否插入额外空行”的选项
    extra_newline_var = tk.BooleanVar()
    tk.Checkbutton(frame, text="段间插入额外空行", variable=extra_newline_var, bg="#f0f0f0").grid(row=5, column=0, columnspan=4, padx=5, pady=5)
    extra_newline_vars.append(extra_newline_var)

    return frame

# 创建选项容器
def create_options_frame(parent, c_start, c_end):
    global current_options_range
    frame = tk.Frame(parent, bg="#f0f0f0")
    canvas = tk.Canvas(frame, bg="#f0f0f0")
    canvas.configure(height=180)
    scrollbar = tk.Scrollbar(frame, orient="horizontal", command=canvas.xview)
    scrollable_frame = tk.Frame(canvas, bg="#f0f0f0")

    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(
            scrollregion=canvas.bbox("all")
        )
    )

    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    canvas.configure(xscrollcommand=scrollbar.set)

    columns = [chr(i) for i in range(ord(c_start), ord(c_end) + 1)]
    for i, col in enumerate(columns):
        tk.Label(scrollable_frame, text=f"列 {col}:", bg="#f0f0f0").grid(row=0, column=i * 3, padx=5, pady=5, sticky="ew")
        bold_header_var = tk.BooleanVar(value=True)
        bold_content_var = tk.BooleanVar()
        colon_var = tk.BooleanVar(value=True)
        newline_var = tk.BooleanVar(value=True)
        tk.Checkbutton(scrollable_frame, text="首行加粗", variable=bold_header_var, bg="#f0f0f0").grid(row=1, column=i * 3, padx=5, pady=5, sticky="w")
        tk.Checkbutton(scrollable_frame, text="内容加粗", variable=bold_content_var, bg="#f0f0f0").grid(row=2, column=i * 3, padx=5, pady=5, sticky="w")
        tk.Checkbutton(scrollable_frame, text="首行后冒号", variable=colon_var, bg="#f0f0f0").grid(row=3, column=i * 3, padx=5, pady=5, sticky="w")
        tk.Checkbutton(scrollable_frame, text="首行后换行", variable=newline_var, bg="#f0f0f0").grid(row=4, column=i * 3, padx=5, pady=5, sticky="w")
        bold_header_vars.append(bold_header_var)
        bold_content_vars.append(bold_content_var)
        colon_vars.append(colon_var)
        newline_vars.append(newline_var)

    canvas.grid(row=0, column=0, sticky="nsew")
    scrollbar.grid(row=1, column=0, sticky="ew", pady=0)

    current_options_range = (c_start, c_end)
    return frame, canvas, scrollable_frame



# 更新选项容器
def update_options_frame():
    # 清空旧的选项取值
    global bold_header_vars, bold_content_vars, colon_vars, newline_vars
    bold_header_vars = []
    bold_content_vars = []
    colon_vars = []
    newline_vars = []
    # 更新选项容器
    global options_frame, canvas, scrollable_frame  # 添加 scrollable_frame 到全局变量
    options_frame.destroy()
    c_start = c_start_var.get()
    c_end = c_end_var.get()
    options_frame, canvas, scrollable_frame = create_options_frame(root, c_start, c_end)
    options_frame.grid(row=2, column=0, sticky="nsew")
    global current_options_range
    current_options_range = (c_start, c_end)


# 创建预览框
def create_preview_frame(parent):
    frame = tk.Frame(parent, bg="#f0f0f0")
    preview_text = scrolledtext.ScrolledText(frame, wrap=tk.WORD, width=40, height=6, bg="white", font=("Arial", 12))
    preview_text.grid(row=0, column=0, padx=5, pady=5, sticky="nsew")
    return frame, preview_text

# 生成预览文本
def generate_preview():
    c_start = c_start_var.get()
    c_end = c_end_var.get()

    # 检查 options_frame 的索引范围是否超出了当前的 c_start 到 c_end 范围
    if current_options_range[0] > c_start or current_options_range[1] < c_end:
        preview_text.delete('1.0', tk.END)
        preview_text.insert(tk.END, "请先点击'更新选项'按钮")
        return

    # 生成预览文本
    preview_text.delete('1.0', tk.END)  # 清空预览框
    paragraph = ""
    for ind in range(ord(c_start), ord(c_end) + 1):
        col = chr(ind)
        i = ind - ord(c_start)
        bold_header = bold_header_vars[i].get()
        bold_content = bold_content_vars[i].get()
        colon = colon_vars[i].get()
        newline = newline_vars[i].get()

        # 添加属性
        if bold_header:
            if colon:
                paragraph += f"第{col}列粗体题目："
            else:
                paragraph += f"第{col}列粗体题目"
        else:
            if colon:
                paragraph += f"第{col}列题目："
            else:
                paragraph += f"第{col}列题目"

        # 添加换行
        if newline:
            paragraph += "\n"

        # 添加内容
        if bold_content:
            paragraph += f"第{col}列粗体内容"
        else:
            paragraph += f"第{col}列内容"

        # 该项属性结束，接着换行
        paragraph += "\n"

    preview_text.insert(tk.END, paragraph + "\n")


# 生成文档
def generate_document():
    excel_file = excel_file_var.get()
    docx_file = docx_file_var.get()
    r_title = r_title_var.get()
    r_start = r_start_var.get()
    r_end = r_end_var.get()
    c_start = c_start_var.get()
    c_end = c_end_var.get()
    extra_line = extra_newline_vars[0].get()

    # 检查 options_frame 的索引范围是否超出了当前的 c_start 到 c_end 范围
    if current_options_range[0] > c_start or current_options_range[1] < c_end:
        preview_text.delete('1.0', tk.END)
        preview_text.insert(tk.END, "请先点击'更新选项'按钮")
        return

    try:

        # 读取Excel文件
        wb = load_workbook(excel_file)
        ws = wb.active

        # 创建Word文档
        doc = Document()

        # 设置字体大小
        style = doc.styles['Normal']
        style.font.size = Pt(14)

        # 获取第一行的单元格内容作为属性
        header_row = ws.iter_rows(min_row=r_title, max_row=r_title, min_col=ord(c_start) - ord('A') + 1, max_col=ord(c_end) - ord('A') + 1)
        headers = [cell.value for cell in next(header_row)]

        # 遍历行和列
        for row in ws.iter_rows(min_row=r_start, max_row=r_end):
            for ind, cell in enumerate(row):
                if ord(c_start) <= ind + ord('A') <= ord(c_end):
                    i = ind - (ord(c_start) - ord('A'))
                    bold_header = bold_header_vars[i].get()
                    bold_content = bold_content_vars[i].get()
                    colon = colon_vars[i].get()
                    newline = newline_vars[i].get()

                    # 添加属性
                    paragraph = doc.add_paragraph()
                    if colon:
                        run = paragraph.add_run(f"{headers[i]}：")
                    else:
                        run = paragraph.add_run(f"{headers[i]}")
                    if bold_header:
                        run.bold = True

                    # 添加换行(插入新段落)
                    if newline:
                        paragraph = doc.add_paragraph()

                    # 添加内容
                    run = paragraph.add_run(str(cell.value))
                    if bold_content:
                        run.bold = True

                    # 该项属性结束，接着换行
                # 该段结束，额外换行
                if extra_line:
                    paragraph = doc.add_paragraph()


        # 保存Word文档
        doc.save(docx_file)

        # 更新预览窗口
        preview_text.delete('1.0', tk.END)
        preview_text.insert(tk.END, "已成功生成文档")

    except Exception as e:
        # 更新预览窗口

        preview_text.delete('1.0', tk.END)
        preview_text.insert(tk.END, f"出错了，未能生成文档: {str(e)}")

# 创建输入框
input_frame = create_input_frame(root)
input_frame.grid(row=0, column=0, padx=5, pady=5, sticky="nw")

update_button = tk.Button(root, text="更新选项", command=update_options_frame, bg="blue", fg="white", font=("Arial", 12))
update_button.grid(row=1, column=0, padx=5, pady=5, sticky="nsew")

# 创建选项容器
options_frame, canvas, scrollable_frame = create_options_frame(root, default_c_start, default_c_end)
options_frame.grid(row=2, column=0, sticky="nw")

# 创建生成预览按钮
preview_button = tk.Button(root, text="生成预览", command=generate_preview, bg="green", fg="white", font=("Arial", 12))
preview_button.grid(row=3, column=0, padx=5, pady=5, sticky="nsew")

# 创建预览框
preview_frame, preview_text = create_preview_frame(root)
preview_frame.grid(row=4, column=0, padx=5, pady=5, sticky="nw")

# 创建生成文档按钮
generate_button = tk.Button(root, text="生成文档", command=generate_document, bg="red", fg="white", font=("Arial", 12))
generate_button.grid(row=5, column=0, padx=5, pady=5, sticky="nsew")

# 配置预览框样式
preview_text.tag_configure('bold', font=('Arial', 12, 'bold'))

# 运行主循环
root.mainloop()
